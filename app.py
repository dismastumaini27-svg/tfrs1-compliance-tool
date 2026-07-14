import streamlit as st
import pandas as pd
from datetime import datetime
import io
import os
import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ------------------------------------------------------------
# CONFIGURATION
# ------------------------------------------------------------
DEPARTMENTS = [
    "Human Resource",
    "Monitoring & Evaluation",
    "Information & Communication Technology",
    "Government Communication Unit",
    "Policy & Planning",
    "Procurement",
    "Commissioner for Minerals",
    "Finance & Accounts",
    "Internal Auditing"
]

YEARS = ["2025", "2026", "2027", "2028", "2029", "2030"]

TFRS_GROUPS = {
    "A. Nature of Operation": "Description of the industry, markets, products, services.",
    "B. Objectives, Strategies & Operating Model": "Strategic objectives, resource allocation, operating model.",
    "C. Resources": "Financial, intellectual, human, social, natural resources.",
    "D. Principal Risks & Uncertainties": "Key risks, likelihood, impact, mitigation.",
    "E. Stakeholder Relationships": "Relationships with employees, customers, suppliers, communities.",
    "F. Capital Structure, Treasury & Liquidity": "Capital structure, treasury policies, cash flows.",
    "G. Performance & KPIs": "Key performance indicators, service delivery targets, efficiency.",
    "H. Corporate Governance": "Governance structure, committees, ethics, auditor appointment.",
    "I. Forward-looking & Future Prospects": "Anticipated trends, challenges, strategic responses.",
    "J. Compliance & Responsibility": "Statement of responsibility, compliance with TFRS 1, publication."
}

# ------------------------------------------------------------
# ALL QUESTIONS – expanded to cover full TFRS 1 requirements
# ------------------------------------------------------------
QUESTION_LIST = [
    # ---------- Human Resource ----------
    ("Human Resource", "A. Nature of Operation",
     "Describe the role of HR in supporting the Ministry's operational mandate, and how is it structured?",
     "Outline staffing levels, organizational structure, and key HR functions (recruitment, payroll, training)."),
    ("Human Resource", "C. Resources",
     "What are the current staff recruitment, retention, and development strategies, and what are the key metrics?",
     "Provide vacancy rates, turnover %, training days per staff, and reference to strategic plans."),
    ("Human Resource", "C. Resources",
     "What employee welfare initiatives (health, pension, safety) are in place, and what is the coverage?",
     "Coverage percentages, budget allocated, specific schemes (PPF/NSSF)."),
    ("Human Resource", "C. Resources",
     "How does the department ensure gender parity and equal opportunity in employment and promotion?",
     "Current gender ratio (M/F). Recruitment/promotion by gender."),
    ("Human Resource", "C. Resources",
     "What training and career progression opportunities are provided for persons with disabilities?",
     "Numbers of disabled staff recruited, trained, promoted."),
    ("Human Resource", "E. Stakeholder Relationships",
     "What grievance handling mechanisms are in place for staff, and what are the resolution statistics?",
     "Number of grievances lodged, resolved, average resolution time."),
    ("Human Resource", "G. Performance & KPIs",
     "What are the key HR performance metrics, and what do the trends show?",
     "Turnover rate, cost per hire, staff satisfaction score, vacancy rate."),
    ("Human Resource", "I. Forward-looking & Future Prospects",
     "What major HR trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "Digital HR, talent shortages, remote work, skills gaps."),

    # ---------- Monitoring & Evaluation ----------
    ("Monitoring & Evaluation", "A. Nature of Operation",
     "How is the M&E function structured, and what is its role in tracking Ministry performance?",
     "Outline the M&E framework, reporting cycles, and key stakeholders."),
    ("Monitoring & Evaluation", "G. Performance & KPIs",
     "What are the actual service delivery targets (outputs) achieved, and what performance indicators were used?",
     "Actual outputs vs. targets (e.g., inspections, reports produced)."),
    ("Monitoring & Evaluation", "G. Performance & KPIs",
     "How efficiently were resources utilized in terms of inputs versus outputs, and what ratios were achieved?",
     "Cost per output, productivity ratios, budget absorption rate."),
    ("Monitoring & Evaluation", "E. Stakeholder Relationships",
     "How has stakeholder feedback and citizen satisfaction been incorporated, and what were the scores?",
     "Satisfaction scores, feedback response rates, actions taken."),
    ("Monitoring & Evaluation", "G. Performance & KPIs",
     "To what extent have the intended societal impacts (outcomes) been achieved, and what is the evidence?",
     "Outcome indicators (e.g., health improvements, economic growth)."),
    ("Monitoring & Evaluation", "H. Corporate Governance",
     "What is the implementation status of previous M&E recommendations (closed, in progress, not implemented)?",
     "Number closed, in progress, not implemented."),
    ("Monitoring & Evaluation", "G. Performance & KPIs",
     "Which Key Performance Indicators (KPIs) are tracked consistently, and what are the trends compared to targets?",
     "Targets vs. actuals for top 5 KPIs."),
    ("Monitoring & Evaluation", "I. Forward-looking & Future Prospects",
     "What major M&E trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "New data collection tools, impact evaluation, real-time reporting."),

    # ---------- ICT ----------
    ("Information & Communication Technology", "A. Nature of Operation",
     "How is the ICT function structured, and what critical services does it provide to the Ministry?",
     "Outline systems supported, user base, and key services (email, networks, databases)."),
    ("Information & Communication Technology", "D. Principal Risks & Uncertainties",
     "What are the key cybersecurity risks, data privacy measures, and IT policies in place?",
     "Number of incidents, compliance level, policy version."),
    ("Information & Communication Technology", "B. Objectives, Strategies & Operating Model",
     "What is the digital transformation strategy, and what progress has been made in implementation?",
     "Milestones achieved (systems launched, automation, adoption rates)."),
    ("Information & Communication Technology", "D. Principal Risks & Uncertainties",
     "What business continuity and disaster recovery plans are in place, and when were they last tested?",
     "Last testing date, Recovery Time Objective (RTO) achieved."),
    ("Information & Communication Technology", "F. Capital Structure, Treasury & Liquidity",
     "What are the costs and benefits of major IT investments, and what is the ROI?",
     "Total capex/opex, ROI, efficiency gains (e.g., time reduced by X%)."),
    ("Information & Communication Technology", "C. Resources",
     "What is the current IT asset inventory, and how is it managed?",
     "Total assets, age profile, replacement value, disposal records."),
    ("Information & Communication Technology", "G. Performance & KPIs",
     "What is the system downtime record, and how quickly are incidents resolved?",
     "Number of outages, average resolution time, root causes."),
    ("Information & Communication Technology", "I. Forward-looking & Future Prospects",
     "What major ICT trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "Cloud migration, AI, mobile apps, data analytics, remote work support."),

    # ---------- Government Communication Unit ----------
    ("Government Communication Unit", "A. Nature of Operation",
     "What is the role of the Government Communication Unit, and what are its core functions?",
     "Outline public engagement, media relations, and information dissemination."),
    ("Government Communication Unit", "B. Objectives, Strategies & Operating Model",
     "What communication strategies and public engagement plans are in place, and what is their reach?",
     "Campaigns, events, media engagements, reach (estimated audience)."),
    ("Government Communication Unit", "G. Performance & KPIs",
     "How effectively does the unit disseminate information to citizens, and what are the metrics?",
     "Number of press releases, website traffic, social media impressions."),
    ("Government Communication Unit", "E. Stakeholder Relationships",
     "What mechanisms are in place for handling citizen feedback/complaints, and how effective are they?",
     "Feedback volumes, response rates, resolution times."),
    ("Government Communication Unit", "J. Compliance & Responsibility",
     "How does the unit ensure transparency and access to information in compliance with the law?",
     "Number of information requests handled, proactive disclosures."),
    ("Government Communication Unit", "G. Performance & KPIs",
     "What public awareness campaigns have been conducted, and what was their effectiveness?",
     "Campaign reach, behavior change results, cost per person reached."),
    ("Government Communication Unit", "D. Principal Risks & Uncertainties",
     "What risks related to misinformation and reputation have been identified, and how are they managed?",
     "Number of reputation incidents, media monitoring outcomes."),
    ("Government Communication Unit", "I. Forward-looking & Future Prospects",
     "What major communication trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "Digital engagement, fake news, citizen journalism, crisis communication."),

    # ---------- Policy & Planning ----------
    ("Policy & Planning", "A. Nature of Operation",
     "What is the role of the Policy & Planning function, and how does it coordinate with national priorities?",
     "Outline policy formulation, planning, and coordination with national priorities."),
    ("Policy & Planning", "B. Objectives, Strategies & Operating Model",
     "How are the Ministry's strategic objectives aligned with national development plans (e.g., Five-Year Plan)?",
     "Map specific goals to national indicators."),
    ("Policy & Planning", "G. Performance & KPIs",
     "What is the progress of strategic plan implementation, and what are the major milestones achieved?",
     "Status of each goal (achieved, ongoing, delayed). Milestones reached."),
    ("Policy & Planning", "A. Nature of Operation",
     "What legislative and regulatory changes affecting the minerals sector have occurred, and what is their impact?",
     "List new laws/amendments and their expected impact on operations."),
    ("Policy & Planning", "B. Objectives, Strategies & Operating Model",
     "What is the policy development framework, and how are stakeholders consulted?",
     "Number of consultations, feedback incorporated."),
    ("Policy & Planning", "D. Principal Risks & Uncertainties",
     "What risks are associated with policy implementation, and what mitigation measures are in place?",
     "Identify top risks and mitigation measures."),
    ("Policy & Planning", "I. Forward-looking & Future Prospects",
     "What major policy trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "New mining codes, environmental regulations, regional integration."),

    # ---------- Procurement ----------
    ("Procurement", "A. Nature of Operation",
     "What is the role of the Procurement function, and what are its critical responsibilities?",
     "Outline procurement of goods/services, contract management, and supplier relations."),
    ("Procurement", "B. Objectives, Strategies & Operating Model",
     "What are the strategic procurement plans and annual budgets, and what is the actual spend?",
     "Total planned vs. actual spend. Number of tenders planned vs issued."),
    ("Procurement", "J. Compliance & Responsibility",
     "What is the level of adherence to the Public Procurement Act regulations, and are there any breaches?",
     "Percentage of tenders fully compliant. Any major breaches/penalties?"),
    ("Procurement", "G. Performance & KPIs",
     "What are the major contracts awarded, and how have suppliers performed?",
     "Top 5 contracts by value. Supplier scores (e.g., 4.5/5)."),
    ("Procurement", "D. Principal Risks & Uncertainties",
     "What risks related to supply chain, fraud, and delays have been identified, and how are they mitigated?",
     "List top risks and controls. Number of fraud incidents."),
    ("Procurement", "G. Performance & KPIs",
     "What value-for-money and cost savings have been achieved through procurement?",
     "Estimated savings from competitive bidding vs. estimated price."),
    ("Procurement", "J. Compliance & Responsibility",
     "How are sustainable and local content procurement policies being implemented?",
     "Percentage local content, environmental criteria applied."),
    ("Procurement", "I. Forward-looking & Future Prospects",
     "What major procurement trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "e-Procurement, strategic sourcing, green procurement."),

    # ---------- Commissioner for Minerals ----------
    ("Commissioner for Minerals", "A. Nature of Operation",
     "What is the mandate of the Commissioner for Minerals, and what are the core regulatory functions?",
     "Outline licensing, inspection, revenue collection, and sector development."),
    ("Commissioner for Minerals", "G. Performance & KPIs",
     "What are the mineral production statistics and royalty collections, and how do they compare to targets?",
     "Tonnes produced (by mineral), total value, royalties collected vs. target."),
    ("Commissioner for Minerals", "A. Nature of Operation",
     "What is the status of licensing processes and monitoring of licensed activities?",
     "Number of new/renewed licenses, compliance rate."),
    ("Commissioner for Minerals", "G. Performance & KPIs",
     "What inspection and compliance activities have been carried out, and what were the findings?",
     "Number of inspections, findings, enforcement actions."),
    ("Commissioner for Minerals", "J. Compliance & Responsibility",
     "What environmental protection and rehabilitation activities have been undertaken?",
     "Rehabilitation area (ha), number of EIAs, closure plans."),
    ("Commissioner for Minerals", "G. Performance & KPIs",
     "How do actual revenue collections compare to targets, and what are the reasons for variances?",
     "Variance analysis: reasons for over/under collection."),
    ("Commissioner for Minerals", "D. Principal Risks & Uncertainties",
     "What risks relate to mining compliance, illegal mining, and revenue leakage, and how are they controlled?",
     "Number of illegal mining cases, revenue lost, controls."),
    ("Commissioner for Minerals", "I. Forward-looking & Future Prospects",
     "What major minerals sector trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "Global commodity prices, new discoveries, local beneficiation policy."),

    # ---------- Finance & Accounts ----------
    ("Finance & Accounts", "F. Capital Structure, Treasury & Liquidity",
     "What is the current cash flow position and liquidity status of the Ministry?",
     "Opening/closing cash balances, cash inflow/outflow, liquidity ratio."),
    ("Finance & Accounts", "G. Performance & KPIs",
     "What is the budget absorption rate, and what are the major variances (actual vs planned)?",
     "Overall budget execution percentage (>90%?). Explain variances >10%."),
    ("Finance & Accounts", "D. Principal Risks & Uncertainties",
     "What financial risks (currency, inflation, fraud) have been identified, and how are they mitigated?",
     "Quantify exposure, risk mitigation measures."),
    ("Finance & Accounts", "F. Capital Structure, Treasury & Liquidity",
     "What treasury policies are in place, and how are public funds managed?",
     "Cash management strategies, bank reconciliation status."),
    ("Finance & Accounts", "F. Capital Structure, Treasury & Liquidity",
     "What significant payments, commitments, and outstanding obligations exist?",
     "Large contracts pending, bills, committed funds."),
    ("Finance & Accounts", "F. Capital Structure, Treasury & Liquidity",
     "What is the financial impact of major capital projects (cost, funding, expected returns)?",
     "Total cost, funding sources, expected returns/benefits."),
    ("Finance & Accounts", "I. Forward-looking & Future Prospects",
     "What major financial trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "Budget constraints, funding reforms, PPPs, revenue diversification."),

    # ---------- Internal Auditing ----------
    ("Internal Auditing", "H. Corporate Governance",
     "What is the annual audit plan and charter, and what is the coverage scope?",
     "Scope coverage, number of planned audits vs completed."),
    ("Internal Auditing", "H. Corporate Governance",
     "What are the key audit findings, recommendations, and management responses?",
     "List top 5 high-risk findings and whether management accepted them."),
    ("Internal Auditing", "H. Corporate Governance",
     "What is the implementation status of previous audit recommendations?",
     "Number closed, in progress, not implemented. Provide percentages."),
    ("Internal Auditing", "H. Corporate Governance",
     "How effective is the internal control environment, and what is the assessment?",
     "Qualitative opinion on control design and operation."),
    ("Internal Auditing", "J. Compliance & Responsibility",
     "What is the compliance status with financial laws and regulations, and were there any breaches?",
     "Breaches detected, penalties, remedial actions."),
    ("Internal Auditing", "D. Principal Risks & Uncertainties",
     "What governance, fraud, and operational risks have been identified, and what is their status?",
     "Risk assessment results, top risks, mitigation status."),
    ("Internal Auditing", "I. Forward-looking & Future Prospects",
     "What major audit trends are anticipated in the next 1-3 years, and what is the strategic response?",
     "Data analytics, continuous auditing, forensic services."),
]

# ---------- GLOBAL QUESTIONS ----------
GLOBAL_QUESTIONS = [
    ("J. Compliance & Responsibility",
     "Has the statement of responsibility (TFRS 1, para 47-48) been formally adopted, confirming accountability for true and fair financial statements?",
     "Draft a formal statement acknowledging responsibility."),
    ("J. Compliance & Responsibility",
     "Has full compliance with TFRS 1 and all relevant laws been declared in the report?",
     "Declare full compliance with TFRS 1 and all relevant laws."),
    ("H. Corporate Governance",
     "Who is the external auditor (name, address, registration, TIN, PF number), and how are they appointed?",
     "Provide full contact details and registration numbers."),
    ("J. Compliance & Responsibility",
     "What political and charitable donations were made, and to which political recipients?",
     "Disclose total political donations and recipients; charitable donations totals (names not required)."),
    ("J. Compliance & Responsibility",
     "Has the report been published on the Ministry website within 30 days of approval?",
     "Confirm if published on the Ministry website within 30 days of approval."),
    ("J. Compliance & Responsibility",
     "Who approved the report, and on what date (list signatories and designations)?",
     "List signatories (names, designations) and approval date."),
    ("H. Corporate Governance",
     "How does the Ministry comply with best practice corporate governance codes?",
     "Explain how the Ministry complies with best practice governance codes."),
    ("H. Corporate Governance",
     "Who are the governance members (Permanent Secretary, Commissioners, etc.), and what is their meeting attendance record?",
     "List governance members and meeting attendance."),
    ("I. Forward-looking & Future Prospects",
     "What are the major strategic priorities for the entire Ministry in the next 1-3 years, and what is the outlook?",
     "Based on national plans, mineral sector strategy, and budget outlook.")
]

# Build the dictionary used by the app
DEPARTMENT_QUESTIONS = {}
for dept, group, q, g in QUESTION_LIST:
    DEPARTMENT_QUESTIONS.setdefault(dept, []).append(
        {"group": group, "question": q, "guidance": g}
    )

# ------------------------------------------------------------
# FILE HELPERS
# ------------------------------------------------------------
DATA_FILE = "tfrs_data"
SYNTHESIS_FILE = "synthesis_data"
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def count_words(text):
    if not text or not isinstance(text, str):
        return 0
    return len(re.findall(r'\b\w+\b', text))

def load_data(year):
    file_key = f"{DATA_FILE}_{year}.csv"
    try:
        if os.path.exists(file_key):
            df = pd.read_csv(file_key)
            required = ['Department', 'Group', 'Question', 'Guidance',
                        'Comments', 'Narrative', 'Attachments', 'Year', 'Last_Updated']
            for col in required:
                if col not in df.columns:
                    df[col] = ''
            df['Year'] = df.get('Year', year)
            return df
        else:
            return create_new_data(year)
    except Exception:
        if os.path.exists(file_key):
            os.remove(file_key)
        return create_new_data(year)

def create_new_data(year):
    rows = []
    for dept, questions in DEPARTMENT_QUESTIONS.items():
        for q in questions:
            rows.append({
                "Department": dept,
                "Group": q["group"],
                "Question": q["question"],
                "Guidance": q["guidance"],
                "Comments": "",
                "Narrative": "",
                "Attachments": "",
                "Year": year,
                "Last_Updated": datetime.now().strftime("%Y-%m-%d %H:%M")
            })
    for group, q, g in GLOBAL_QUESTIONS:
        rows.append({
            "Department": "**CORPORATE / GENERAL**",
            "Group": group,
            "Question": q,
            "Guidance": g,
            "Comments": "",
            "Narrative": "",
            "Attachments": "",
            "Year": year,
            "Last_Updated": datetime.now().strftime("%Y-%m-%d %H:%M")
        })
    df = pd.DataFrame(rows)
    df.to_csv(f"{DATA_FILE}_{year}.csv", index=False)
    return df

def save_data(df, year):
    df.to_csv(f"{DATA_FILE}_{year}.csv", index=False)

def load_synthesis(year):
    file_key = f"{SYNTHESIS_FILE}_{year}.csv"
    try:
        if os.path.exists(file_key):
            df = pd.read_csv(file_key)
            if 'Group' not in df.columns or 'Synthesis' not in df.columns:
                os.remove(file_key)
                raise Exception("Bad format")
            for group in TFRS_GROUPS.keys():
                if group not in df['Group'].values:
                    df = pd.concat([df, pd.DataFrame([{'Group': group, 'Synthesis': ''}])], ignore_index=True)
            return df[['Group', 'Synthesis']]
        else:
            rows = [{'Group': group, 'Synthesis': ''} for group in TFRS_GROUPS.keys()]
            df = pd.DataFrame(rows)
            df.to_csv(file_key, index=False)
            return df
    except Exception:
        if os.path.exists(file_key):
            os.remove(file_key)
        rows = [{'Group': group, 'Synthesis': ''} for group in TFRS_GROUPS.keys()]
        df = pd.DataFrame(rows)
        df.to_csv(file_key, index=False)
        return df

def save_synthesis(df, year):
    df = df[['Group', 'Synthesis']]
    df.to_csv(f"{SYNTHESIS_FILE}_{year}.csv", index=False)

# ------------------------------------------------------------
# STREAMLIT APP
# ------------------------------------------------------------
st.set_page_config(layout="wide")
st.title("🏛 TFRS 1 Report Builder – Ministry of Minerals")

try:
    # Sidebar
    st.sidebar.title("📋 Reporting Period")
    selected_year = st.sidebar.selectbox("Select Financial Year", YEARS)

    st.sidebar.markdown("---")
    st.sidebar.title("📌 Department")
    selected_dept = st.sidebar.selectbox("Select Your Department",
                                         DEPARTMENTS + ["**CORPORATE / GENERAL**"])

    st.sidebar.markdown("---")
    st.sidebar.caption("📌 Write at least **100 words** per narrative.")
    st.sidebar.caption("📌 Attach supporting files.")

    # Load data for selected year
    if 'data' not in st.session_state or st.session_state.get('current_year') != selected_year:
        st.session_state.data = load_data(selected_year)
        st.session_state.current_year = selected_year
        st.session_state.synthesis = load_synthesis(selected_year)

    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = {}

    # Admin: Add Question
    with st.sidebar.expander("🔧 Admin - Add New Question (Password: admin123)"):
        admin_pass = st.text_input("Password", type="password", key="admin_pass")
        if admin_pass == "admin123":
            st.success("Admin access granted.")
            new_dept = st.selectbox("Department", DEPARTMENTS + ["**CORPORATE / GENERAL**"])
            new_group = st.selectbox("TFRS Group", list(TFRS_GROUPS.keys()))
            new_q = st.text_area("Question")
            new_guidance = st.text_area("Guidance")
            if st.button("➕ Add Question"):
                if new_q and new_guidance:
                    new_row = {
                        "Department": new_dept,
                        "Group": new_group,
                        "Question": new_q,
                        "Guidance": new_guidance,
                        "Comments": "",
                        "Narrative": "",
                        "Attachments": "",
                        "Year": selected_year,
                        "Last_Updated": datetime.now().strftime("%Y-%m-%d %H:%M")
                    }
                    st.session_state.data = pd.concat([st.session_state.data, pd.DataFrame([new_row])], ignore_index=True)
                    save_data(st.session_state.data, selected_year)
                    st.success("✅ Added!")
                    st.rerun()
        elif admin_pass:
            st.error("Wrong password.")

    st.sidebar.markdown("---")
    st.sidebar.caption(f"Financial Year {selected_year}")

    # ---------- DATA ENTRY ----------
    st.subheader(f"📋 Checklist: {selected_dept}")
    st.caption(f"**Reporting Period:** Financial Year {selected_year}")
    st.caption("Minimum 100 words per narrative.")

    dept_mask = st.session_state.data["Department"] == selected_dept
    dept_data = st.session_state.data[dept_mask].copy()

    if dept_data.empty:
        st.warning(f"No questions found for {selected_dept}.")
        st.stop()

    with st.form(key="entry_form"):
        updated_narratives = []
        updated_comments = []
        updated_attachments = []

        q_counter = 1
        for idx, row in dept_data.iterrows():
            question = row["Question"]
            guidance = row["Guidance"]
            current_narrative = row["Narrative"]
            current_comment = row["Comments"]
            current_attachments = row["Attachments"]

            with st.expander(f"Q{q_counter}: {question}", expanded=False):
                col1, col2 = st.columns([2, 1])
                with col1:
                    st.caption(f"💡 {guidance}")
                    narrative = st.text_area(
                        "Narrative",
                        value=current_narrative,
                        key=f"narrative_{idx}",
                        height=150,
                        placeholder="Write at least 100 words..."
                    )
                    comment = st.text_input(
                        "Comments / References",
                        value=current_comment,
                        key=f"comment_{idx}",
                        placeholder="e.g., see policy page 5"
                    )
                with col2:
                    wc = count_words(narrative)
                    if wc >= 100:
                        st.success(f"✅ {wc} words - Compliant!")
                    else:
                        st.warning(f"⚠️ {wc} / 100 words")
                        st.progress(wc / 100)

                    uploaded_file = st.file_uploader(
                        "Attach document",
                        type=["pdf", "xlsx", "xls", "docx", "png", "jpg", "jpeg"],
                        key=f"upload_{idx}",
                        label_visibility="collapsed"
                    )
                    if uploaded_file is not None:
                        ext = uploaded_file.name.split('.')[-1]
                        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                        safe_name = f"{selected_dept.replace(' ', '_')}_{ts}.{ext}"
                        file_path = os.path.join(UPLOAD_DIR, safe_name)
                        with open(file_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())
                        st.success(f"✅ File saved: {safe_name}")
                        current_attachments = safe_name if not current_attachments else current_attachments + ", " + safe_name
                    if current_attachments:
                        st.info(f"📎 Attached: {current_attachments}")

                updated_narratives.append(narrative)
                updated_comments.append(comment)
                updated_attachments.append(current_attachments)

            q_counter += 1

        submitted = st.form_submit_button("💾 Save My Data")
        if submitted:
            for i, (idx, row) in enumerate(dept_data.iterrows()):
                st.session_state.data.at[idx, "Narrative"] = updated_narratives[i]
                st.session_state.data.at[idx, "Comments"] = updated_comments[i]
                st.session_state.data.at[idx, "Attachments"] = updated_attachments[i]
                st.session_state.data.at[idx, "Last_Updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.session_state.data.at[idx, "Year"] = selected_year
            save_data(st.session_state.data, selected_year)
            st.success("✅ Data saved successfully!")

    # ---------- DASHBOARD ----------
    st.markdown("---")
    st.header("📊 Live Compliance Dashboard")
    st.caption(f"**Reporting Period:** Financial Year {selected_year}")
    st.info("💡 **Compliance %** = (questions with ≥100 words) / (total questions).")

    all_depts = st.session_state.data["Department"].unique()
    summary = []
    for dept in all_depts:
        if dept == "**CORPORATE / GENERAL**":
            continue
        dept_filter = st.session_state.data[st.session_state.data["Department"] == dept]
        total = len(dept_filter)
        compliant = sum(1 for _, r in dept_filter.iterrows() if count_words(r["Narrative"]) >= 100)
        compliance = (compliant / total * 100) if total > 0 else 0
        summary.append({
            "Department": dept,
            "Total Questions": total,
            "✅ Compliant": compliant,
            "❌ Incomplete": total - compliant,
            "Compliance %": round(compliance, 1)
        })
    df_summary = pd.DataFrame(summary)
    st.dataframe(df_summary, use_container_width=True, hide_index=True)
    st.subheader("Overall Compliance by Department (%)")
    st.bar_chart(df_summary.set_index("Department")["Compliance %"])

    # ---------- ADMIN SYNTHESIS ----------
    st.markdown("---")
    st.header("🔧 Admin Consolidation / Synthesis")

    with st.expander("📝 Admin: Synthesize Departmental Inputs (Password: admin123)"):
        synth_pass = st.text_input("Enter Admin Password", type="password", key="synth_pass")
        if synth_pass == "admin123":
            st.success("✅ Admin access granted.")
            st.info("Rewrite raw inputs into one cohesive paragraph per TFRS group.")

            st.subheader("📄 Raw Departmental Inputs")
            for group in TFRS_GROUPS.keys():
                group_data = st.session_state.data[
                    (st.session_state.data["Group"] == group) &
                    (st.session_state.data["Department"] != "**CORPORATE / GENERAL**")
                ]
                if not group_data.empty:
                    with st.expander(f"View Raw Data for {group}"):
                        for dept in DEPARTMENTS:
                            dept_group = group_data[group_data["Department"] == dept]
                            if not dept_group.empty:
                                st.markdown(f"**{dept}:**")
                                for _, row in dept_group.iterrows():
                                    st.caption(f"Q: {row['Question']}")
                                    st.write(row['Narrative'] if row['Narrative'] else "⚠ No data provided.")
                                    st.caption(f"Word count: {count_words(row['Narrative'])}")
                                st.divider()

            st.subheader("✍️ Write Your Polished Synthesis Paragraphs")
            synth_df = st.session_state.synthesis.copy()
            updated = {}
            for group in TFRS_GROUPS.keys():
                current = synth_df[synth_df['Group'] == group]['Synthesis'].values[0] if not synth_df[synth_df['Group'] == group].empty else ""
                new_text = st.text_area(
                    f"Synthesis for {group}",
                    value=current,
                    height=120,
                    key=f"synth_{group}",
                    placeholder="Write a cohesive paragraph here..."
                )
                updated[group] = new_text

            if st.button("💾 Save All Synthesis Text"):
                for group, text in updated.items():
                    idx = st.session_state.synthesis[st.session_state.synthesis['Group'] == group].index
                    if not idx.empty:
                        st.session_state.synthesis.at[idx[0], 'Synthesis'] = text
                    else:
                        st.session_state.synthesis = pd.concat([st.session_state.synthesis, pd.DataFrame([{'Group': group, 'Synthesis': text}])], ignore_index=True)
                save_synthesis(st.session_state.synthesis, selected_year)
                st.success(f"✅ Synthesis saved for FY {selected_year}!")
                st.rerun()
        elif synth_pass:
            st.error("Wrong password.")

    # ---------- REPORT GENERATOR ----------
    st.markdown("---")
    st.header("📄 Generate Draft TFRS 1 Report")
    st.warning(f"Report for Financial Year {selected_year}.")

    if st.button("📝 Generate Consolidated Draft Report (Word)"):
        output = io.BytesIO()
        if DOCX_AVAILABLE:
            doc = Document()
            doc.add_heading('REPORT BY THOSE CHARGED WITH GOVERNANCE', 0)
            doc.add_heading('Ministry of Minerals', level=1)
            doc.add_paragraph(f'Reporting Period: Financial Year {selected_year}')
            doc.add_paragraph(f'Date: {datetime.now().strftime("%d %B %Y")}')
            doc.add_paragraph('Prepared in accordance with TFRS 1')
            doc.add_paragraph('')

            doc.add_heading('1. Executive Summary', level=2)
            for _, row in df_summary.iterrows():
                doc.add_paragraph(f'• {row["Department"]}: {row["Compliance %"]}% compliance ({row["✅ Compliant"]} out of {row["Total Questions"]})')
            doc.add_paragraph('')

            corp_data = st.session_state.data[st.session_state.data["Department"] == "**CORPORATE / GENERAL**"]
            if not corp_data.empty:
                doc.add_heading('2. Corporate Governance', level=2)
                for _, row in corp_data.iterrows():
                    doc.add_heading(row["Question"], level=3)
                    doc.add_paragraph(row["Narrative"] if row["Narrative"] else "⚠ Not provided.")
                    if row["Attachments"]:
                        doc.add_paragraph(f'Attachments: {row["Attachments"]}')
                doc.add_paragraph('')

            doc.add_heading('3. Operational and Financial Review by TFRS Section', level=2)
            synth_df = st.session_state.synthesis
            for group in TFRS_GROUPS.keys():
                doc.add_heading(f'3.{list(TFRS_GROUPS.keys()).index(group)+1} {group}', level=3)
                doc.add_paragraph(TFRS_GROUPS[group])
                synth_row = synth_df[synth_df['Group'] == group]
                if not synth_row.empty and synth_row.iloc[0]['Synthesis'] and synth_row.iloc[0]['Synthesis'].strip():
                    doc.add_paragraph(synth_row.iloc[0]['Synthesis'], style='List Bullet')
                    doc.add_paragraph("*(Consolidated by Management.)*", style='List Bullet')
                else:
                    group_data = st.session_state.data[
                        (st.session_state.data["Group"] == group) &
                        (st.session_state.data["Department"] != "**CORPORATE / GENERAL**")
                    ]
                    if group_data.empty:
                        doc.add_paragraph('⚠ No data provided.')
                    else:
                        for dept in DEPARTMENTS:
                            dept_group = group_data[group_data["Department"] == dept]
                            if dept_group.empty:
                                continue
                            doc.add_heading(f'{dept}:', level=4)
                            for _, row in dept_group.iterrows():
                                doc.add_paragraph(f'• {row["Question"]}')
                                if row["Narrative"]:
                                    doc.add_paragraph(row["Narrative"], style='List Bullet')
                                else:
                                    doc.add_paragraph('⚠ Not provided.', style='List Bullet')
                                if row["Attachments"]:
                                    doc.add_paragraph(f'Evidence: {row["Attachments"]}', style='List Bullet')
                doc.add_paragraph('')

            doc.save(output)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            output.write("="*80 + "\n".encode())
            output.write("REPORT BY THOSE CHARGED WITH GOVERNANCE - MINISTRY OF MINERALS\n".encode())
            output.write(f"Financial Year {selected_year}\n".encode())
            output.write("="*80 + "\n\n".encode())
            output.write("(Install python-docx for a proper Word document.)\n".encode())
            mime = "text/plain"
            ext = "txt"

        st.download_button(
            label=f"⬇ Download Draft Report (.{ext})",
            data=output.getvalue(),
            file_name=f"TFRS1_FY{selected_year}_{datetime.now().strftime('%Y%m%d')}.{ext}",
            mime=mime
        )
        st.success(f"Report generated for FY {selected_year}!")

except Exception as e:
    st.error(f"An error occurred: {e}")
    st.stop()
